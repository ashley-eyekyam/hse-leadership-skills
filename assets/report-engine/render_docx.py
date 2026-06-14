"""
render_docx.py — ported EyekyamDOCX renderer (python-docx), theme-token driven.

PORT of ~/.claude/skills/eyekyam-docx/scripts/create_docx.py's EyekyamDOCX
class. The public method set mirrors render_pdf.EyekyamPDF (add_cover,
add_heading, add_text, add_section, add_bullets, add_callout, add_code_block,
add_table, add_metrics_row, add_two_columns, add_image, add_divider,
add_page_break, add_spacer, add_contact_section, add_toc_placeholder, build) —
that matched API IS the dispatch contract the shared generate_report.py driver
(Plan 02-06) keys off.

NOTE (matched-API asymmetry, by design): EyekyamDOCX has add_callout but NOT
add_info_box (the PDF renderer has both). The shared dispatch maps the
`callout` block type to add_callout on both renderers, so they stay matched.

Exactly three surgical edits were applied to the port source:

  Edit 1 (font seam): the module-level DEFAULT_FONT = "Calibri" is replaced by
    a theme-driven body font (theme.fonts.body, default "Noto Sans"). The OXML
    helpers take an explicit font_name. python-docx embeds by font NAME
    (name-only sufficient for v1.0; embedTrueTypeFonts deferred per A4 §4.10
    assumption A3). The bundled Noto family carries ₹ and (named) Devanagari.

  Edit 2 (brand/company seam): the module-level `class BRAND` / `class COMPANY`
    constant reads are replaced by values pulled from an injected Theme on the
    EyekyamDOCX constructor (theme.palette.* / theme.org.* / theme.images.logo).

  Edit 3 (logo-on-teal): _logo_on_white_card() composites the logo onto a white
    card so it stays legible on the teal cover hero, with the Pillow import
    guarded — absent Pillow degrades to the raw logo (A4 §4.10).
"""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Iterable, List, Optional, Sequence, Tuple, Union

from docx import Document
from docx.document import Document as _DocxDoc
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

from theme import resolve_theme


# ---------------------------------------------------------------------------
# Module-level fallback font name (Edit 1). Instance methods pass the resolved
# theme body font explicitly; this is only the degradation default.
# ---------------------------------------------------------------------------
_FALLBACK_FONT = "Noto Sans"


# ---------------------------------------------------------------------------
# Low-level XML helpers — unchanged from the port source (OXML is subtle and
# already correct). _add_page_number_field now takes an explicit font_name.
# ---------------------------------------------------------------------------
def _shading_element(hex_color: str) -> OxmlElement:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    return shd


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(_shading_element(hex_color))


def _shade_paragraph(paragraph, hex_color: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    pPr.append(_shading_element(hex_color))


def _set_cell_borders(cell, color: str = "DDDDDD", sz: str = "4") -> None:
    """Apply a thin border on all four sides of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_cell_no_borders(cell) -> None:
    """Strip all visible borders from a cell (useful for cover hero)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _mark_repeat_header(row) -> None:
    """Make this row repeat as a header on every page the table spans."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(OxmlElement("w:tblHeader"))


def _add_page_number_field(paragraph, color_hex: Optional[str] = None,
                           size_pt: int = 8, font_name: str = _FALLBACK_FONT) -> None:
    """Insert a Word PAGE field (auto-updating page number)."""
    run = paragraph.add_run()
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _add_horizontal_rule(paragraph, color_hex: str = "2AADA8",
                         size: str = "12") -> None:
    """Bottom border on a paragraph — visually a horizontal rule."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    for old in pBdr.findall(qn("w:bottom")):
        pBdr.remove(old)
    pBdr.append(bottom)


# ---------------------------------------------------------------------------
# Inline-formatting parser — unchanged from the port source. _add_runs takes
# an explicit font_name (Edit 1) so the resolved theme body font is honoured.
# ---------------------------------------------------------------------------
_TAG_RE = re.compile(r"(<[^>]+>)")
_ENTITY_MAP = {
    "&amp;": "\x01AMP\x01",
    "&lt;": "<",
    "&gt;": ">",
    "&nbsp;": " ",
}


def _unescape_entities(text: str) -> str:
    for needle, repl in _ENTITY_MAP.items():
        text = text.replace(needle, repl)
    return text.replace("\x01AMP\x01", "&")


def _parse_inline(text: str):
    """Yield (segment, attrs) tuples for each run of text."""
    text = _unescape_entities(text)
    parts = _TAG_RE.split(text)
    bold = italic = underline = False
    color: Optional[str] = None
    for chunk in parts:
        if not chunk:
            continue
        if chunk.startswith("<") and chunk.endswith(">"):
            inner = chunk[1:-1].strip()
            closing = inner.startswith("/")
            tag = inner.lstrip("/").split()[0].lower() if inner else ""
            if tag == "b":
                bold = not closing
            elif tag == "i":
                italic = not closing
            elif tag == "u":
                underline = not closing
            elif tag.startswith("br"):
                yield ("\n", {"br": True})
            elif tag == "font":
                if closing:
                    color = None
                else:
                    m = re.search(r'color\s*=\s*"([^"]+)"', inner, re.IGNORECASE)
                    if m:
                        color = m.group(1).lstrip("#")
        else:
            yield (chunk, {
                "bold": bold,
                "italic": italic,
                "underline": underline,
                "color": color,
            })


def _add_runs(paragraph, text: str, *,
              base_color: Optional[str] = None,
              base_size: Optional[float] = None,
              base_bold: bool = False,
              font_name: str = _FALLBACK_FONT) -> None:
    """Append text to ``paragraph`` honouring inline HTML formatting."""
    if text is None:
        return
    for segment, attrs in _parse_inline(str(text)):
        if attrs.get("br"):
            paragraph.add_run().add_break()
            continue
        run = paragraph.add_run(segment)
        run.font.name = font_name
        if base_size is not None:
            run.font.size = Pt(base_size)
        run.bold = bool(attrs.get("bold")) or base_bold
        run.italic = bool(attrs.get("italic"))
        run.underline = bool(attrs.get("underline"))
        rgb = attrs.get("color") or base_color
        if rgb:
            run.font.color.rgb = RGBColor.from_string(rgb.lstrip("#"))


# ---------------------------------------------------------------------------
# Page geometry — A4 with the same margins the PDF renderer uses
# ---------------------------------------------------------------------------
PAGE_W_CM = 21.0
PAGE_H_CM = 29.7
MARGIN_CM = 2.0
USABLE_W_CM = PAGE_W_CM - 2 * MARGIN_CM   # 17.0
LETTER_W_CM = 21.59
LETTER_H_CM = 27.94


# ===========================================================================
# Main class
# ===========================================================================
class EyekyamDOCX:
    """
    Mirrors the EyekyamPDF API, driven by a resolved Theme (theme.py).
    Method names/params match the PDF renderer so the shared dispatch driver
    walks either target identically.
    """

    def __init__(self, output_path: str, title: str,
                 doc_type: str = "report",
                 author: str = "Eyekyam Risk Resolutions",
                 subject: str = "",
                 theme=None) -> None:
        self.output_path = output_path
        self.title = title
        self.doc_type = doc_type
        self.author = author
        self.subject = subject
        self.theme = theme if theme is not None else resolve_theme()

        # ── Edit 1: body font from the resolved theme ──
        self.DEFAULT_FONT = self.theme.fonts.body or _FALLBACK_FONT
        # ── Edit 2: palette (hex w/o '#') + org + logo from the theme ──
        self._build_brand()

        # Page geometry honours theme layout page_size.
        if str(self.theme.layout.get("page_size")) == "Letter":
            self._page_w_cm, self._page_h_cm = LETTER_W_CM, LETTER_H_CM
        else:
            self._page_w_cm, self._page_h_cm = PAGE_W_CM, PAGE_H_CM
        self._usable_w_cm = self._page_w_cm - 2 * MARGIN_CM

        self.doc: _DocxDoc = Document()
        cp = self.doc.core_properties
        cp.title = title
        cp.author = author
        if subject:
            cp.subject = subject

        style = self.doc.styles["Normal"]
        style.font.name = self.DEFAULT_FONT
        style.font.size = Pt(10)

        section = self.doc.sections[0]
        section.page_width = Cm(self._page_w_cm)
        section.page_height = Cm(self._page_h_cm)
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

        self._strip_default_paragraph()

        self._cover_section_added = False
        self._body_section: Optional[object] = None
        self._running_chrome_installed = False
        self._logo_white_bg_path: Optional[str] = None

    # ── Edit 2: brand tokens (hex without '#') from the theme palette ──
    def _build_brand(self):
        p = self.theme.palette

        def hx(key, fallback):
            return str(p.get(key, fallback)).lstrip("#").upper()

        risk = p.get("risk", {})

        class _Brand:
            PRIMARY_TEAL = hx("primary", "2AADA8")
            DEEP_TEAL = hx("secondary", "1A7070")
            LIGHT_TEAL = hx("accent", "47C4BE")
            PALE_TEAL = "E8F7F7"
            NEAR_BLACK = hx("text", "1A1A1A")
            WHITE = "FFFFFF"
            LIGHT_GRAY = hx("surface", "F5F5F5")
            MID_GRAY = "888888"
            BORDER_LIGHT = "DDDDDD"
            HIGH_RISK = str(risk.get("high", "E05252")).lstrip("#").upper()
            MED_RISK = str(risk.get("med", "E9A94B")).lstrip("#").upper()
            LOW_RISK = str(risk.get("low", "4CAF50")).lstrip("#").upper()
            CRITICAL_RISK = str(risk.get("critical", "B71C1C")).lstrip("#").upper()

        self.BRAND = _Brand
        org = self.theme.org or {}
        self._org_name = org.get("name") or ""
        self._org_footer = org.get("footer") or ""
        self._classification = org.get("classification") or "Confidential"
        self._logo_path = self.theme.images.get("logo")

    def _strip_default_paragraph(self) -> None:
        body = self.doc.element.body
        for child in list(body):
            if child.tag == qn("w:p") and not child.text and len(child) == 0:
                body.remove(child)
                break

    # ── Edit 3: logo-on-white-card (Pillow-guarded) ──
    def _logo_on_white_card(self) -> Optional[str]:
        """Return a path to the logo composited onto a white card so it reads
        on the teal cover. Falls back to the raw logo if Pillow is absent or
        compositing fails — never raises."""
        if not self._logo_path:
            return None
        if self._logo_white_bg_path and Path(self._logo_white_bg_path).exists():
            return self._logo_white_bg_path
        try:
            from PIL import Image
            src = Image.open(self._logo_path).convert("RGBA")
            pad = max(40, src.width // 25)
            bg = Image.new("RGBA",
                           (src.width + 2 * pad, src.height + 2 * pad),
                           (255, 255, 255, 255))
            bg.paste(src, (pad, pad), src)
            out_path = Path(self._logo_path).parent / "eyekyam-logo-white-bg.png"
            bg.convert("RGB").save(out_path, "PNG")
            self._logo_white_bg_path = str(out_path)
            return self._logo_white_bg_path
        except Exception:
            return self._logo_path

    def set_logo(self, path: str) -> None:
        """Override the resolved logo (instance-scoped)."""
        self._logo_path = str(path)
        self._logo_white_bg_path = None

    # ==================================================================
    # Cover page
    # ==================================================================
    def add_cover(self, title: str = "", subtitle: str = "",
                  client: str = "", date: str = "",
                  tagline: str = "") -> None:
        """Render a full-bleed teal cover page with logo + title block."""
        BRAND = self.BRAND
        DEFAULT_FONT = self.DEFAULT_FONT
        page_w_cm, page_h_cm = self._page_w_cm, self._page_h_cm

        cover_section = self.doc.sections[0]
        cover_section.top_margin = Cm(0)
        cover_section.bottom_margin = Cm(0)
        cover_section.left_margin = Cm(0)
        cover_section.right_margin = Cm(0)
        cover_section.header_distance = Cm(0)
        cover_section.footer_distance = Cm(0)
        cover_section.different_first_page_header_footer = True
        cover_section.first_page_header.is_linked_to_previous = False
        cover_section.first_page_footer.is_linked_to_previous = False
        for para in cover_section.first_page_header.paragraphs:
            para.clear()
        for para in cover_section.first_page_footer.paragraphs:
            para.clear()

        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Cm(page_w_cm)
        cell = table.cell(0, 0)
        cell.width = Cm(page_w_cm)
        row = table.rows[0]
        row.height = Cm(page_h_cm)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _shade_cell(cell, BRAND.PRIMARY_TEAL)
        _set_cell_no_borders(cell)
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for edge, val in (("top", "0"), ("left", "720"), ("bottom", "0"), ("right", "720")):
            m = OxmlElement(f"w:{edge}")
            m.set(qn("w:w"), val)
            m.set(qn("w:type"), "dxa")
            tcMar.append(m)
        tcPr.append(tcMar)

        first_para = cell.paragraphs[0]
        first_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        first_para.paragraph_format.space_before = Pt(36)

        # ── Edit 3: logo on white card (Pillow-guarded) ──
        logo_card = self._logo_on_white_card()
        if logo_card and Path(logo_card).exists():
            run = first_para.add_run()
            try:
                run.add_picture(logo_card, width=Cm(4.5))
            except Exception:
                pass

        sp = cell.add_paragraph()
        sp.paragraph_format.space_before = Pt(120)

        if title:
            for line in title.split("\n"):
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Cm(2.0)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(line)
                run.font.name = DEFAULT_FONT
                run.font.size = Pt(34)
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(BRAND.WHITE)

        if subtitle:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(2.0)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(subtitle)
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor.from_string(BRAND.PALE_TEAL)

        if client:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(2.0)
            p.paragraph_format.space_before = Pt(28)
            run = p.add_run(client)
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(BRAND.WHITE)

        if date:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(2.0)
            p.paragraph_format.space_before = Pt(4)
            run = p.add_run(date)
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string(BRAND.PALE_TEAL)

        if tagline:
            spacer = cell.add_paragraph()
            spacer.paragraph_format.space_before = Pt(200)
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(2.0)
            run = p.add_run(tagline)
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(10)
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(BRAND.PALE_TEAL)

        body_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        body_section.page_width = Cm(page_w_cm)
        body_section.page_height = Cm(page_h_cm)
        body_section.top_margin = Cm(MARGIN_CM)
        body_section.bottom_margin = Cm(MARGIN_CM)
        body_section.left_margin = Cm(MARGIN_CM)
        body_section.right_margin = Cm(MARGIN_CM)
        body_section.header_distance = Cm(1.0)
        body_section.footer_distance = Cm(1.0)
        body_section.different_first_page_header_footer = False
        self._body_section = body_section
        self._cover_section_added = True
        self._install_running_chrome(body_section)

    # ── Running header + footer ──
    def _install_running_chrome(self, section) -> None:
        if self._running_chrome_installed:
            return
        self._running_chrome_installed = True
        BRAND = self.BRAND
        DEFAULT_FONT = self.DEFAULT_FONT

        header = section.header
        h_para = header.paragraphs[0]
        h_para.clear()
        h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _shade_paragraph(h_para, BRAND.PRIMARY_TEAL)
        h_para.paragraph_format.space_before = Pt(2)
        h_para.paragraph_format.space_after = Pt(2)
        run = h_para.add_run(self.title)
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BRAND.WHITE)

        footer = section.footer
        f_para = footer.paragraphs[0]
        f_para.clear()
        f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade_paragraph(f_para, BRAND.DEEP_TEAL)
        f_para.paragraph_format.space_before = Pt(2)
        f_para.paragraph_format.space_after = Pt(2)
        from datetime import datetime
        year = datetime.now().year
        owner = self._org_name or "HSE Report Engine"
        footer_bits = self._org_footer or ""
        prefix = f"© {year} {owner}"
        if footer_bits:
            prefix += f" | {footer_bits}"
        prefix += "        Page "
        run = f_para.add_run(prefix)
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(BRAND.WHITE)
        _add_page_number_field(f_para, color_hex=BRAND.WHITE, size_pt=8,
                               font_name=DEFAULT_FONT)

    def _ensure_chrome(self) -> None:
        if self._running_chrome_installed:
            return
        section = self.doc.sections[0]
        self._install_running_chrome(section)

    # ==================================================================
    # Headings, body text, callouts
    # ==================================================================
    def add_heading(self, text: str, level: int = 1) -> None:
        BRAND = self.BRAND
        font = self.DEFAULT_FONT
        self._ensure_chrome()
        p = self.doc.add_paragraph()
        if level == 1:
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            _add_runs(p, text, base_color=BRAND.PRIMARY_TEAL,
                      base_size=20, base_bold=True, font_name=font)
            p.paragraph_format.keep_with_next = True
            _add_horizontal_rule(p, BRAND.PRIMARY_TEAL, size="6")
        elif level == 2:
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            _add_runs(p, text, base_color=BRAND.DEEP_TEAL,
                      base_size=14, base_bold=True, font_name=font)
            p.paragraph_format.keep_with_next = True
        else:
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            _add_runs(p, text, base_color=BRAND.NEAR_BLACK,
                      base_size=11, base_bold=True, font_name=font)
            p.paragraph_format.keep_with_next = True

    def add_text(self, text: str) -> None:
        """Justified body paragraph supporting <b>, <i>, <u>, <br/>, <font>."""
        self._ensure_chrome()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        _add_runs(p, text, base_color=self.BRAND.NEAR_BLACK, base_size=10,
                  font_name=self.DEFAULT_FONT)

    def add_section(self, heading: str, body: str, level: int = 1) -> None:
        self.add_heading(heading, level=level)
        self.add_text(body)

    def add_bullets(self, items: Iterable[str], symbol: str = "•") -> None:
        BRAND = self.BRAND
        font = self.DEFAULT_FONT
        self._ensure_chrome()
        for item in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(f"{symbol}  ")
            r.font.name = font
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor.from_string(BRAND.NEAR_BLACK)
            _add_runs(p, item, base_color=BRAND.NEAR_BLACK, base_size=10,
                      font_name=font)

    def add_callout(self, text: str) -> None:
        """Light-teal info box for key insights."""
        BRAND = self.BRAND
        self._ensure_chrome()
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.columns[0].width = Cm(self._usable_w_cm)
        cell = table.cell(0, 0)
        cell.width = Cm(self._usable_w_cm)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade_cell(cell, BRAND.PALE_TEAL)
        _set_cell_borders(cell, color=BRAND.LIGHT_TEAL, sz="6")
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for edge in ("top", "left", "bottom", "right"):
            m = OxmlElement(f"w:{edge}")
            m.set(qn("w:w"), "180")
            m.set(qn("w:type"), "dxa")
            tcMar.append(m)
        tcPr.append(tcMar)
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        _add_runs(para, text, base_color=BRAND.DEEP_TEAL, base_size=10,
                  font_name=self.DEFAULT_FONT)
        self.add_spacer(0.3)

    def add_code_block(self, lines: Sequence[str], font_size: int = 9,
                       title: Optional[str] = None) -> None:
        """Monospace code block inside a light-grey-shaded cell."""
        BRAND = self.BRAND
        mono = self.theme.fonts.mono or "Courier New"
        if title:
            self.add_heading(title, level=3)
        self._ensure_chrome()
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.columns[0].width = Cm(self._usable_w_cm)
        cell = table.cell(0, 0)
        cell.width = Cm(self._usable_w_cm)
        _shade_cell(cell, BRAND.LIGHT_GRAY)
        _set_cell_borders(cell, color=BRAND.BORDER_LIGHT, sz="4")
        cell.paragraphs[0].clear()
        for i, line in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line if line else " ")
            run.font.name = mono
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor.from_string(BRAND.NEAR_BLACK)
        self.add_spacer(0.3)

    # ==================================================================
    # Tables
    # ==================================================================
    def add_table(self, headers: List[str], rows: List[List[str]],
                  col_widths: Optional[List[float]] = None,
                  caption: Optional[str] = None,
                  risk_col: Optional[int] = None) -> None:
        """Add a branded data table (params mirror the PDF renderer)."""
        BRAND = self.BRAND
        font = self.DEFAULT_FONT
        self._ensure_chrome()
        n = len(headers)
        POINTS_PER_CM = 28.346456692913385
        if not col_widths:
            widths_cm = [self._usable_w_cm / n] * n
        elif max(col_widths) >= 25:
            widths_cm = [w / POINTS_PER_CM for w in col_widths]
        else:
            widths_cm = list(col_widths)
        total = sum(widths_cm)
        if total > self._usable_w_cm:
            scale = self._usable_w_cm / total
            widths_cm = [w * scale for w in widths_cm]

        table = self.doc.add_table(rows=1 + len(rows), cols=n)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        for col_idx, w in enumerate(widths_cm):
            for row_cells in table.rows:
                row_cells.cells[col_idx].width = Cm(w)

        header_row = table.rows[0]
        _mark_repeat_header(header_row)
        for col_idx, h in enumerate(headers):
            cell = header_row.cells[col_idx]
            _shade_cell(cell, BRAND.PRIMARY_TEAL)
            _set_cell_borders(cell, color=BRAND.DEEP_TEAL, sz="6")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _add_runs(p, h, base_color=BRAND.WHITE, base_size=9,
                      base_bold=True, font_name=font)

        risk_map = {
            "high": BRAND.HIGH_RISK,
            "medium": BRAND.MED_RISK,
            "med": BRAND.MED_RISK,
            "low": BRAND.LOW_RISK,
            "critical": BRAND.CRITICAL_RISK,
        }
        for r_idx, row_data in enumerate(rows):
            zebra = BRAND.WHITE if r_idx % 2 == 0 else BRAND.LIGHT_GRAY
            row_cells = table.rows[r_idx + 1].cells
            for c_idx, val in enumerate(row_data):
                cell = row_cells[c_idx]
                _shade_cell(cell, zebra)
                _set_cell_borders(cell, color=BRAND.BORDER_LIGHT, sz="4")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                color = BRAND.NEAR_BLACK
                bold = False
                if risk_col is not None and c_idx == risk_col:
                    risk_color = risk_map.get(str(val).strip().lower())
                    if risk_color:
                        color = risk_color
                        bold = True
                _add_runs(p, str(val), base_color=color, base_size=9,
                          base_bold=bold, font_name=font)

        if caption:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(2)
            cap.paragraph_format.space_after = Pt(6)
            run = cap.add_run(caption)
            run.font.name = font
            run.font.size = Pt(8)
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(BRAND.MID_GRAY)

        self.add_spacer(0.3)

    # ==================================================================
    # KPI metrics row
    # ==================================================================
    def add_metrics_row(self, metrics: List[dict]) -> None:
        """Render 2-4 KPI boxes side by side."""
        BRAND = self.BRAND
        font = self.DEFAULT_FONT
        self._ensure_chrome()
        n = len(metrics)
        if n == 0:
            return
        col_w = self._usable_w_cm / n
        table = self.doc.add_table(rows=1, cols=n)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        for i, m in enumerate(metrics):
            cell = table.rows[0].cells[i]
            cell.width = Cm(col_w)
            color = (m.get("color") or BRAND.PRIMARY_TEAL).lstrip("#")
            _shade_cell(cell, color)
            _set_cell_borders(cell, color=color, sz="4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_val = cell.paragraphs[0]
            p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_val.paragraph_format.space_before = Pt(8)
            p_val.paragraph_format.space_after = Pt(0)
            run = p_val.add_run(str(m.get("value", "")))
            run.font.name = font
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(BRAND.WHITE)
            unit = m.get("unit")
            if unit:
                u_run = p_val.add_run(f" {unit}")
                u_run.font.name = font
                u_run.font.size = Pt(11)
                u_run.font.color.rgb = RGBColor.from_string(BRAND.WHITE)
            p_lbl = cell.add_paragraph()
            p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_lbl.paragraph_format.space_before = Pt(2)
            p_lbl.paragraph_format.space_after = Pt(8)
            run = p_lbl.add_run(str(m.get("label", "")))
            run.font.name = font
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(BRAND.WHITE)
        self.add_spacer(0.4)

    # ==================================================================
    # Two-column layout
    # ==================================================================
    def add_two_columns(self, left, right, left_w_pct: float = 0.5) -> None:
        """Side-by-side layout."""
        self._ensure_chrome()
        left_w = self._usable_w_cm * left_w_pct
        right_w = self._usable_w_cm - left_w
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.rows[0].cells[0].width = Cm(left_w)
        table.rows[0].cells[1].width = Cm(right_w)
        for cell, content in zip(table.rows[0].cells, [left, right]):
            _set_cell_no_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            items = content if isinstance(content, (list, tuple)) else [content]
            cell.paragraphs[0].clear()
            for i, txt in enumerate(items):
                p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                _add_runs(p, str(txt), base_color=self.BRAND.NEAR_BLACK,
                          base_size=10, font_name=self.DEFAULT_FONT)
        self.add_spacer(0.3)

    # ==================================================================
    # Image, divider, page break, spacer
    # ==================================================================
    def add_image(self, path: str, width_cm: float = 14.0,
                  caption: Optional[str] = None,
                  align: str = "CENTER") -> None:
        """Embed a PNG/JPEG image with optional caption."""
        BRAND = self.BRAND
        self._ensure_chrome()
        align_map = {
            "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
            "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
            "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        p = self.doc.add_paragraph()
        p.alignment = align_map.get(align.upper(), WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
        if caption:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(6)
            r = cap.add_run(caption)
            r.font.name = self.DEFAULT_FONT
            r.font.size = Pt(8)
            r.italic = True
            r.font.color.rgb = RGBColor.from_string(BRAND.MID_GRAY)

    def add_divider(self) -> None:
        """Horizontal rule."""
        self._ensure_chrome()
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        _add_horizontal_rule(p, self.BRAND.BORDER_LIGHT, size="6")

    def add_page_break(self) -> None:
        """Force a new page."""
        self._ensure_chrome()
        p = self.doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

    def add_spacer(self, height_cm: float = 0.5) -> None:
        """Vertical whitespace via an empty paragraph with space-before."""
        self._ensure_chrome()
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(max(1, int(height_cm * 28.35)))
        p.paragraph_format.space_after = Pt(0)

    # ==================================================================
    # Contact + TOC
    # ==================================================================
    def add_contact_section(self, heading: str = "Contact") -> None:
        """Branded contact block based on the theme org tokens."""
        self.add_heading(heading, level=2)
        org = self.theme.org or {}
        lines = []
        if org.get("name"):
            lines.append(f"<b>{org['name']}</b>")
        if org.get("footer"):
            lines.append(org["footer"])
        if not lines:
            lines = ["—"]
        for ln in lines:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _add_runs(p, ln, base_color=self.BRAND.NEAR_BLACK, base_size=10,
                      font_name=self.DEFAULT_FONT)

    def add_toc_placeholder(self, sections: list,
                            auto_number: Optional[bool] = None) -> None:
        """Render a Table of Contents and follow with a page break."""
        font = self.DEFAULT_FONT
        _NUMBERED_RE = re.compile(r"^\s*\d+(\.\d+)*\.?\s")
        items: List[Tuple[str, str]] = []
        for s in sections:
            if isinstance(s, (tuple, list)) and len(s) >= 2:
                items.append((s[0], s[1]))
            else:
                items.append(("section", str(s)))
        if auto_number is None:
            section_titles = [t for k, t in items if k == "section"] or \
                             [t for _, t in items]
            if section_titles:
                pre_numbered = sum(
                    1 for t in section_titles if _NUMBERED_RE.match(t))
                auto_number = pre_numbered < max(1, len(section_titles) / 2)
            else:
                auto_number = True
        self.add_heading("Table of Contents", level=2)
        section_counter = 1
        for kind, title in items:
            text = str(title)
            if kind == "sub":
                indent = Cm(0.8)
                italic_open, italic_close = "", ""
                bold_wrap = False
            elif kind == "phase":
                indent = Cm(0.8)
                italic_open, italic_close = "<i>", "</i>"
                bold_wrap = False
            else:
                indent = Cm(0.0)
                italic_open, italic_close = "", ""
                bold_wrap = True
            if auto_number and kind == "section":
                prefix = f"<b>{section_counter}.</b>  "
                section_counter += 1
            else:
                prefix = ""
            wrapped = (f"<b>{text}</b>" if bold_wrap and not prefix else text)
            line = f"{prefix}{italic_open}{wrapped}{italic_close}"
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = indent
            p.paragraph_format.space_after = Pt(3)
            _add_runs(p, line, base_color=self.BRAND.NEAR_BLACK, base_size=10,
                      font_name=font)
        self.add_page_break()

    # ==================================================================
    # Build / save
    # ==================================================================
    def build(self) -> str:
        """Save the .docx and return the output path."""
        self._ensure_chrome()
        self.doc.save(self.output_path)
        return self.output_path
